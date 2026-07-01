"""Word (.docx) 文档解析，提取段落、表格与嵌入图片。"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, List

from rag.ingestion.asset_store import AssetStore
from rag.ingestion.models import ImageBlock, ParsedDocument, TextBlock


def parse_docx(
    file_path: Path,
    *,
    doc_id: str,
    display_path: str,
    asset_store: AssetStore,
) -> ParsedDocument:
    """解析 DOCX 为段落、表格文本与嵌入图片块。

    功能:
        按文档 body 顺序遍历段落与表格；提取文本、表格转「表格：」格式；
        从 drawing/blip 关系提取图片 blob 并写入资产目录。

    参数:
        file_path: DOCX 文件路径。
        doc_id: 文档标识符。
        display_path: 展示用路径。
        asset_store: 图片资产存储。

    返回值:
        含有序 blocks 的 ``ParsedDocument``。

    异常:
        RuntimeError: 未安装 python-docx 依赖。
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise RuntimeError("未安装 python-docx，请执行 pip install python-docx") from exc

    document = Document(str(file_path))
    blocks: List = []
    order = 0
    image_counter = 0

    for item in _iter_body_items(document):
        if isinstance(item, Paragraph):
            text = (item.text or "").strip()
            if text:
                blocks.append(TextBlock(text=text, order=order))
                order += 1
            for rel_id in _iter_paragraph_image_rel_ids(item):
                image_counter += 1
                rel = _save_related_image(
                    document,
                    rel_id,
                    doc_id=doc_id,
                    asset_store=asset_store,
                    image_counter=image_counter,
                )
                if rel:
                    blocks.append(
                        ImageBlock(
                            asset_id=f"media_{image_counter:03d}",
                            relative_path=rel,
                            order=order,
                        )
                    )
                    order += 1
        elif isinstance(item, Table):
            table_text = _table_to_text(item)
            if table_text:
                blocks.append(TextBlock(text=table_text, order=order))
                order += 1
            for rel_id in _iter_table_image_rel_ids(item):
                image_counter += 1
                rel = _save_related_image(
                    document,
                    rel_id,
                    doc_id=doc_id,
                    asset_store=asset_store,
                    image_counter=image_counter,
                )
                if rel:
                    blocks.append(
                        ImageBlock(
                            asset_id=f"media_{image_counter:03d}",
                            relative_path=rel,
                            order=order,
                        )
                    )
                    order += 1

    return ParsedDocument(
        doc_id=doc_id,
        source_path=str(file_path),
        display_path=display_path,
        blocks=blocks,
    )


def _iter_body_items(document) -> Iterable[object]:
    """按文档顺序 yield body 下的段落与表格对象。

    功能:
        遍历 document.element.body 子节点，``p`` 映射为 Paragraph，``tbl`` 为 Table。

    参数:
        document: python-docx Document 实例。

    返回值:
        Paragraph 或 Table 对象的生成器。

    异常:
        无。
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield Paragraph(child, document)
        elif tag == "tbl":
            yield Table(child, document)


def _iter_paragraph_image_rel_ids(paragraph) -> Iterable[str]:
    """从段落 XML 中提取嵌入/链接图片的关系 ID。

    功能:
        在 w:drawing 下查找 a:blip 的 r:embed 或 r:link 属性。

    参数:
        paragraph: python-docx Paragraph 实例。

    返回值:
        关系 ID 字符串的生成器。

    异常:
        无。
    """
    from docx.oxml.ns import qn

    for drawing in paragraph._element.xpath(".//w:drawing"):
        for blip in drawing.xpath(".//a:blip"):
            rel_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if rel_id:
                yield rel_id


def _iter_table_image_rel_ids(table) -> Iterable[str]:
    """从表格所有单元格段落中收集图片关系 ID。

    功能:
        遍历 table 每行每单元格每段落，委托 ``_iter_paragraph_image_rel_ids``。

    参数:
        table: python-docx Table 实例。

    返回值:
        关系 ID 字符串的生成器。

    异常:
        无。
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield from _iter_paragraph_image_rel_ids(paragraph)


def _save_related_image(
    document,
    rel_id: str,
    *,
    doc_id: str,
    asset_store: AssetStore,
    image_counter: int,
) -> str:
    """将 DOCX 关系部件中的图片 blob 保存至资产目录。

    功能:
        从 document.part.related_parts 读取 blob，按 partname 推断后缀并 save_bytes。

    参数:
        document: python-docx Document 实例。
        rel_id: 图片关系 ID。
        doc_id: 文档标识符。
        asset_store: 资产存储。
        image_counter: 当前图片序号，用于生成 asset_id。

    返回值:
        相对 persist 的路径；部件缺失或 blob 为空时返回 ``""``。

    异常:
        无。
    """
    part = document.part.related_parts.get(rel_id)
    if part is None:
        return ""
    blob = getattr(part, "blob", b"") or b""
    if not blob:
        return ""
    partname = str(getattr(part, "partname", ""))
    suffix = Path(partname).suffix.lower() or ".png"
    asset_id = f"media_{image_counter:03d}"
    return asset_store.save_bytes(doc_id, asset_id, blob, suffix=suffix)


def _table_to_text(table) -> str:
    """将 Word 表格转换为 pipe 分隔的纯文本。

    功能:
        每行单元格文本以 `` | `` 连接，多行前缀 ``表格：\\n``。

    参数:
        table: python-docx Table 实例。

    返回值:
        表格文本；空表返回 ``""``。

    异常:
        无。
    """
    lines: List[str] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            cells.append(text)
        line = " | ".join(cells).strip()
        if line:
            lines.append(line)
    if not lines:
        return ""
    return "表格：\n" + "\n".join(lines)
